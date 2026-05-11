"""Shared utilities for the resume-tailor pipeline."""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Any


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
REFERENCES_DIR = SKILL_ROOT / "references"


# ---------- JSON I/O ----------

def load_json(path: str | Path) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_json(obj: Any, path: str | Path, *, indent: int = 2) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=indent, ensure_ascii=False)
        f.write("\n")


# ---------- Input ingest (tier-based) ----------

def read_text_input(path_or_text: str) -> str:
    """Read text from a file path or treat the argument as inline text.

    Supported file extensions: .txt, .md, .pdf, .docx
    Returns plain text. For inline text, returns the argument unchanged.
    """
    p = Path(path_or_text)
    if not p.exists() or not p.is_file():
        # treat as inline text
        return path_or_text

    suffix = p.suffix.lower()
    if suffix in {".txt", ".md"}:
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        return _read_pdf(p)
    if suffix == ".docx":
        return _read_docx(p)
    # unknown extension -> read as text
    return p.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> str:
    try:
        import fitz  # pymupdf
    except ImportError as e:
        raise RuntimeError(
            "Reading PDFs requires pymupdf. Install: pip install pymupdf"
        ) from e
    parts = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "Reading .docx requires python-docx. Install: pip install python-docx"
        ) from e
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # also capture tables (some resumes use them)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------- Synonyms ----------

def load_synonyms(path: Path | None = None) -> dict[str, str]:
    """Load synonyms.json and flatten to a lowercase alias->canonical map.

    Empty alias values are skipped. The leading-underscore meta keys are skipped.
    """
    syn_path = path or (REFERENCES_DIR / "synonyms.json")
    if not syn_path.exists():
        return {}
    raw = load_json(syn_path)
    flat: dict[str, str] = {}
    for category, entries in raw.items():
        if category.startswith("_") or not isinstance(entries, dict):
            continue
        for canonical, aliases in entries.items():
            if canonical.startswith("_"):
                continue
            flat[canonical.lower()] = canonical
            if isinstance(aliases, list):
                for a in aliases:
                    flat[a.lower()] = canonical
    return flat


def canonicalize(token: str, synonyms: dict[str, str] | None = None) -> str:
    """Return the canonical form of a token if known, otherwise the token unchanged."""
    syn = synonyms if synonyms is not None else load_synonyms()
    return syn.get(token.lower().strip(), token.strip())


# ---------- Tokenization ----------

# Strip Markdown / decorative punctuation but keep tech-significant chars (#+./)
_TOKEN_STRIP = re.compile(r"[^\w#+.\-/]")


def tokenize(text: str) -> list[str]:
    """Tokenize text into lowercase tokens, preserving #/+/-/./ for tech terms (C#, C++, .NET).

    Strips trailing/leading punctuation so 'Kafka.' tokenizes as 'kafka' (not 'kafka.'),
    preserving internal punctuation in tokens like 'node.js' and 'C++'.
    """
    cleaned = _TOKEN_STRIP.sub(" ", text)
    out: list[str] = []
    for raw in cleaned.split():
        t = raw.strip(".,;:!?-").lower()
        if t:
            out.append(t)
    return out


# ---------- Number extraction ----------

# Matches: 35%, $14K, 2.4×, 8M, 5+, 12 services, 4-engineer, 200K req/day
_NUMBER_RE = re.compile(
    r"""
    (?<![A-Za-z])           # not preceded by a letter
    (?:\$|US\$)?            # optional currency
    \d+(?:[.,]\d+)?         # int or decimal
    \s*
    (?:[KkMmBb]|×|x(?=\b)|%|\+|/[A-Za-z]+)?   # SI suffix, multiplier, percent, plus
    """,
    re.VERBOSE,
)


def extract_numbers(text: str) -> list[str]:
    """Extract human-readable numeric tokens from a string ('35%', '$14K', '2.4×', '8M')."""
    out = []
    for m in _NUMBER_RE.finditer(text):
        s = m.group(0).strip()
        if s and any(c.isdigit() for c in s):
            out.append(_normalize_number(s))
    return out


def _normalize_number(s: str) -> str:
    """Normalize a numeric token for comparison: lowercase, strip whitespace, unify × vs x."""
    s = s.replace(" ", "").lower()
    s = s.replace("x", "×") if s.endswith("x") else s
    return s


# ---------- CLI ergonomics ----------

def err(*args: Any) -> None:
    print(*args, file=sys.stderr)
