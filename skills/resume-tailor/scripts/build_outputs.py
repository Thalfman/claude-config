"""Render outputs from the tailored markdown resume.

Usage:
    python -m scripts.build_outputs tailored_resume.md jd_analysis.json relevance.json \
        [--formats md,docx,pdf] [--output-dir outputs/]

Produces:
    <First>_<Last>_Resume_<Role>_<Company>.md
    <First>_<Last>_Resume_<Role>_<Company>.docx   (if 'docx' in formats)
    <First>_<Last>_Resume_<Role>_<Company>.pdf    (if 'pdf' in formats; falls back to instruction)
    mapping_report.md
    coverage_scorecard.md

DOCX output uses an ATS-safe template: single-column, no tables in body, standard
fonts, plain bullets, dates in 'MMM YYYY – Present' format.
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from datetime import date
from pathlib import Path
from typing import Any

from ._common import err, load_json


SAFE_NAME = re.compile(r"[^A-Za-z0-9]+")


def slugify(s: str | None, default: str = "X") -> str:
    if not s:
        return default
    out = SAFE_NAME.sub("_", s).strip("_")
    return out or default


def derive_filename(tailored_md: str, jd: dict[str, Any]) -> str:
    """Build <First>_<Last>_Resume_<Role>_<Company> from tailored md + jd analysis."""
    first = "Candidate"
    last = ""
    # H1 first line is conventionally the candidate's name
    for ln in tailored_md.splitlines():
        s = ln.strip()
        if s.startswith("# "):
            name_parts = s[2:].strip().split()
            if name_parts:
                first = name_parts[0]
                last = "_".join(name_parts[1:]) if len(name_parts) > 1 else ""
            break
    role = slugify((jd.get("role_title") or "Role").split(",")[0].split(" at ")[0])
    company = slugify(_guess_company(jd))
    today = date.today().strftime("%Y%m%d")
    parts = [first, last, "Resume", role, company, today]
    return "_".join(p for p in parts if p)


def _guess_company(jd: dict[str, Any]) -> str:
    """Look for 'at <Company>' or 'with <Company>' patterns in the role title."""
    title = jd.get("role_title") or ""
    m = re.search(r"\bat\s+([A-Z][A-Za-z0-9 &\-]+)", title)
    if m:
        return m.group(1).strip()
    return "Company"


# ---------- Markdown copy ----------

def copy_markdown(tailored_md: str, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(tailored_md, encoding="utf-8")


# ---------- DOCX render ----------

def render_docx(tailored_md: str, output_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError(
            "DOCX output requires python-docx. Install: pip install python-docx"
        ) from e

    doc = Document()

    # ATS-safe defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    lines = tailored_md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        s = line.strip()

        # Skip HTML comments
        if s.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # Headings
        if s.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(s[2:].strip())
            run.bold = True
            run.font.size = Pt(20)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif s.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(s[3:].strip().upper())
            run.bold = True
            run.font.size = Pt(13)
        elif s.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(s[4:].strip())
            run.bold = True
            run.font.size = Pt(12)
        elif s.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, s[2:].strip())
        elif s == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, s)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _add_inline_runs(p, text: str) -> None:
    """Render Markdown bold/italic as Word runs."""
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = p.add_run(token[2:-2])
            run.bold = True
        else:
            run = p.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


# ---------- PDF render ----------

def render_pdf(docx_path: Path, output_path: Path) -> bool:
    """Try docx2pdf; fall back to printing instructions for manual conversion."""
    try:
        from docx2pdf import convert  # type: ignore
    except ImportError:
        err(f"docx2pdf not installed. To produce PDF, install it (pip install docx2pdf) "
            f"or open {docx_path.name} in Word and Save As PDF.")
        return False
    try:
        convert(str(docx_path), str(output_path))
        return True
    except Exception as e:  # noqa: BLE001
        err(f"docx2pdf failed: {e}. Open {docx_path.name} in Word and Save As PDF.")
        return False


# ---------- Reports ----------

def render_mapping_report(relevance: dict[str, Any]) -> str:
    lines: list[str] = ["# Mapping Report", ""]
    lines.append("Each JD requirement and the inventory bullet(s) used to address it.")
    lines.append("")
    for r in relevance.get("requirements", []):
        coverage_marker = {"STRONG": "✓", "PARTIAL": "~", "GAP": "✗"}[r["coverage"]]
        lines.append(f"## [{coverage_marker} {r['coverage']}] {r['requirement']}")
        lines.append(f"_kind: {r['kind']}, top_score: {r['top_score']}_")
        lines.append("")
        if r["evidence"]:
            for ev in r["evidence"][:3]:
                lines.append(f"- **{ev['company']}** ({ev['title']}) [score {ev['score']}]")
                lines.append(f"  > {ev['bullet_text']}")
                if ev.get("reasons"):
                    lines.append(f"  - Match reasons: {'; '.join(ev['reasons'])}")
                lines.append("")
        else:
            lines.append("- _No matching evidence found._")
            lines.append("")
    return "\n".join(lines) + "\n"


def render_coverage_scorecard(jd: dict[str, Any], relevance: dict[str, Any], tailored_md: str) -> str:
    s = relevance.get("summary", {})
    lines: list[str] = ["# Coverage Scorecard", ""]
    lines.append(f"**Role:** {jd.get('role_title') or 'Unknown'}")
    lines.append(f"**Level:** {jd.get('level') or 'Unknown'}")
    lines.append("")
    lines.append("## Required-qualification coverage")
    lines.append("")
    lines.append(f"- Strong: {s.get('required_strong', 0)} / {s.get('required_total', 0)}")
    lines.append(f"- Partial: {s.get('required_partial', 0)} / {s.get('required_total', 0)}")
    lines.append(f"- Gap: {s.get('required_gap', 0)} / {s.get('required_total', 0)}")
    lines.append(f"- **Weighted coverage: {s.get('required_coverage_pct', 0)}%**")
    lines.append("")

    # ATS keyword density
    keywords = jd.get("keywords") or []
    tailored_lower = tailored_md.lower()
    word_count = max(1, len(tailored_lower.split()))
    lines.append("## ATS keyword density (in tailored resume)")
    lines.append("")
    lines.append("| Keyword | JD mentions | Resume mentions | Density |")
    lines.append("|---|---:|---:|---:|")
    for kw in keywords[:25]:
        token = kw["keyword"]
        # case-insensitive count, word boundary unless special chars
        if any(c in token for c in "+#./"):
            n = tailored_lower.count(token.lower())
        else:
            n = len(re.findall(rf"\b{re.escape(token.lower())}\b", tailored_lower))
        density = round(100 * n / word_count, 2)
        lines.append(f"| {token} | {kw['count']} | {n} | {density}% |")
    lines.append("")

    # Gaps
    gap_reqs = [r for r in relevance.get("requirements", []) if r["coverage"] == "GAP"]
    if gap_reqs:
        lines.append("## Gaps (no inventory match)")
        lines.append("")
        for r in gap_reqs:
            lines.append(f"- _{r['kind']}_: {r['requirement']}")
        lines.append("")

    return "\n".join(lines) + "\n"


# ---------- CLI ----------

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Build resume outputs from tailored markdown.")
    ap.add_argument("tailored", help="Path to tailored resume (.md)")
    ap.add_argument("jd", help="Path to jd_analysis.json")
    ap.add_argument("relevance", help="Path to relevance.json")
    ap.add_argument("--formats", default="md,docx", help="Comma-separated: md,docx,pdf")
    ap.add_argument("--output-dir", "-o", default="outputs", help="Output directory")
    args = ap.parse_args(argv)

    tailored_md = Path(args.tailored).read_text(encoding="utf-8")
    jd = load_json(args.jd)
    relevance = load_json(args.relevance)

    formats = {f.strip().lower() for f in args.formats.split(",") if f.strip()}
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    base = derive_filename(tailored_md, jd)
    md_path = out_dir / f"{base}.md"
    docx_path = out_dir / f"{base}.docx"
    pdf_path = out_dir / f"{base}.pdf"

    written: list[Path] = []

    if "md" in formats:
        copy_markdown(tailored_md, md_path)
        written.append(md_path)

    if "docx" in formats:
        render_docx(tailored_md, docx_path)
        written.append(docx_path)

    if "pdf" in formats:
        if not docx_path.exists():
            render_docx(tailored_md, docx_path)
            written.append(docx_path)
        if render_pdf(docx_path, pdf_path):
            written.append(pdf_path)

    mapping_path = out_dir / "mapping_report.md"
    mapping_path.write_text(render_mapping_report(relevance), encoding="utf-8")
    written.append(mapping_path)

    scorecard_path = out_dir / "coverage_scorecard.md"
    scorecard_path.write_text(render_coverage_scorecard(jd, relevance, tailored_md), encoding="utf-8")
    written.append(scorecard_path)

    print(f"Wrote {len(written)} files to {out_dir}/:")
    for p in written:
        print(f"  - {p.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
